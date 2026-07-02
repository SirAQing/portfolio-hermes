"""轻量文档解析器 — 支持 Markdown/TXT/HTML/JSON/PDF/DOCX

设计原则：
- 不复制 docreader（避免重依赖）
- 支持格式：.md / .txt / .html / .json / .pdf / .docx
- 统一输出 Markdown 纯文本（供 chunker 使用）
- 可扩展：新格式只需注册 Parser
"""
import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Callable


@dataclass
class ParseResult:
    """解析结果"""
    content: str  # Markdown 纯文本
    title: str = ""
    source_type: str = ""
    meta: dict = None

    def __post_init__(self):
        if self.meta is None:
            self.meta = {}


class ParseError(Exception):
    pass


# ── 解析器注册 ──

_PARSERS: dict[str, Callable[[bytes, str], ParseResult]] = {}


def register_parser(*extensions: str):
    """装饰器：注册解析器。

    extensions: 文件扩展名（不含点），如 "md", "txt", "html"
    """
    def decorator(func: Callable[[bytes, str], ParseResult]):
        for ext in extensions:
            _PARSERS[ext.lower()] = func
        return func
    return decorator


def get_supported_extensions() -> list[str]:
    """返回支持的文件扩展名列表。"""
    return sorted(_PARSERS.keys())


def parse_file(filename: str, content: bytes) -> ParseResult:
    """解析文件，返回 ParseResult。

    filename: 原始文件名（用于推断类型）
    content: 文件原始字节
    """
    ext = Path(filename).suffix.lstrip(".").lower()
    if ext not in _PARSERS:
        raise ParseError(
            f"Unsupported file type: .{ext}. Supported: {get_supported_extensions()}"
        )
    parser = _PARSERS[ext]
    return parser(content, filename)


# ── Markdown 解析器 ──

@register_parser("md", "markdown")
def parse_markdown(content: bytes, filename: str) -> ParseResult:
    """Markdown 直接返回（已是目标格式）。"""
    text = content.decode("utf-8", errors="replace")
    # 提取首个 H1 作为标题
    title = ""
    h1_match = re.search(r"^#\s+(.+)$", text, re.MULTILINE)
    if h1_match:
        title = h1_match.group(1).strip()
    else:
        # 用文件名
        title = Path(filename).stem
    return ParseResult(
        content=text,
        title=title,
        source_type="markdown",
        meta={"char_count": len(text)},
    )


# ── TXT 解析器 ──

@register_parser("txt", "text")
def parse_text(content: bytes, filename: str) -> ParseResult:
    """纯文本直接返回。"""
    # 尝试多种编码
    for encoding in ["utf-8", "gbk", "latin-1"]:
        try:
            text = content.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        text = content.decode("utf-8", errors="replace")

    title = Path(filename).stem
    return ParseResult(
        content=text,
        title=title,
        source_type="text",
        meta={"char_count": len(text), "encoding": encoding},
    )


# ── HTML 解析器（轻量，不依赖 BeautifulSoup） ──

@register_parser("html", "htm")
def parse_html(content: bytes, filename: str) -> ParseResult:
    """简单 HTML → Markdown 转换。

    不依赖 BeautifulSoup，用正则做轻量转换：
    - 移除 script/style/meta 标签
    - h1-h6 → Markdown 标题
    - p/div → 段落
    - a → [text](href)
    - ul/ol/li → 列表
    - code/pre → 代码块
    """
    html = content.decode("utf-8", errors="replace")

    # 提取 title
    title_match = re.search(r"<title>(.*?)</title>", html, re.IGNORECASE | re.DOTALL)
    title = title_match.group(1).strip() if title_match else Path(filename).stem

    # 移除 script/style/meta/head
    html = re.sub(r"<(script|style|meta|link|head)[^>]*>.*?</\1>", "", html, flags=re.IGNORECASE | re.DOTALL)
    # 移除 HTML 注释
    html = re.sub(r"<!--.*?-->", "", html, flags=re.DOTALL)

    # 转换标题
    for i in range(1, 7):
        html = re.sub(
            rf"<h{i}[^>]*>(.*?)</h{i}>",
            lambda m, level=i: f"\n{'#' * level} {m.group(1).strip()}\n",
            html, flags=re.IGNORECASE | re.DOTALL,
        )

    # 转换链接
    html = re.sub(
        r"<a[^>]+href=[\"']([^\"']+)[\"'][^>]*>(.*?)</a>",
        r"[\2](\1)",
        html, flags=re.IGNORECASE | re.DOTALL,
    )

    # 转换列表
    html = re.sub(r"<li[^>]*>(.*?)</li>", r"- \1\n", html, flags=re.IGNORECASE | re.DOTALL)
    html = re.sub(r"</?(ul|ol)[^>]*>", "\n", html, flags=re.IGNORECASE)

    # 转换代码块
    html = re.sub(r"<pre[^>]*>(.*?)</pre>", lambda m: f"\n```\n{m.group(1).strip()}\n```\n", html, flags=re.IGNORECASE | re.DOTALL)
    html = re.sub(r"<code[^>]*>(.*?)</code>", r"`\1`", html, flags=re.IGNORECASE | re.DOTALL)

    # 段落/换行
    html = re.sub(r"<br\s*/?>", "\n", html, flags=re.IGNORECASE)
    html = re.sub(r"</p>", "\n\n", html, flags=re.IGNORECASE)
    html = re.sub(r"</div>", "\n", html, flags=re.IGNORECASE)

    # 移除所有剩余标签
    text = re.sub(r"<[^>]+>", "", html)

    # 清理 HTML 实体
    entities = {
        "&amp;": "&", "&lt;": "<", "&gt;": ">", "&quot;": '"',
        "&apos;": "'", "&nbsp;": " ", "&#39;": "'",
    }
    for ent, char in entities.items():
        text = text.replace(ent, char)
    # 数字实体
    text = re.sub(r"&#(\d+);", lambda m: chr(int(m.group(1))), text)

    # 压缩多余空行
    text = re.sub(r"\n{3,}", "\n\n", text).strip()

    return ParseResult(
        content=text,
        title=title,
        source_type="html",
        meta={"char_count": len(text), "original_format": "html"},
    )


# ── JSON 解析器（结构化文档） ──

@register_parser("json")
def parse_json(content: bytes, filename: str) -> ParseResult:
    """JSON 文档 → Markdown。

    支持两种格式：
    1. {title, content} — 直接提取
    2. 任意 JSON — 序列化为可读文本
    """
    try:
        data = json.loads(content.decode("utf-8", errors="replace"))
    except json.JSONDecodeError as e:
        raise ParseError(f"Invalid JSON: {e}")

    title = Path(filename).stem

    if isinstance(data, dict):
        # 常见字段提取
        if "title" in data:
            title = str(data["title"])
        if "content" in data:
            content_str = f"# {title}\n\n{data['content']}"
        elif "text" in data:
            content_str = f"# {title}\n\n{data['text']}"
        elif "body" in data:
            content_str = f"# {title}\n\n{data['body']}"
        else:
            # 序列化整个对象
            content_str = f"# {title}\n\n```json\n{json.dumps(data, indent=2, ensure_ascii=False)}\n```"
    elif isinstance(data, list):
        # 数组 → 多段
        parts = [f"# {title}\n"]
        for i, item in enumerate(data, 1):
            if isinstance(item, dict):
                parts.append(f"## Item {i}\n")
                for k, v in item.items():
                    parts.append(f"**{k}**: {v}\n")
            else:
                parts.append(f"## Item {i}\n{item}\n")
        content_str = "\n".join(parts)
    else:
        content_str = f"# {title}\n\n{str(data)}"

    return ParseResult(
        content=content_str,
        title=title,
        source_type="json",
        meta={"char_count": len(content_str)},
    )


# ── PDF 解析器 ──

@register_parser("pdf")
def parse_pdf(content: bytes, filename: str) -> ParseResult:
    """PDF 文档 → Markdown 纯文本。

    使用 PyMuPDF (fitz) 提取文本。
    """
    try:
        import fitz
    except ImportError:
        raise ParseError(
            "PDF support requires PyMuPDF. Install with: pip install PyMuPDF"
        )

    if not content or len(content) < 100:
        raise ParseError("PDF file is empty or too small")

    try:
        doc = None
        try:
            doc = fitz.open(stream=content, filetype="pdf")
        except Exception:
            try:
                from io import BytesIO
                doc = fitz.open(stream=BytesIO(content), filetype="pdf")
            except Exception:
                doc = fitz.open(filename=filename, stream=content)

        if doc.is_encrypted:
            try:
                doc.authenticate("")
            except Exception:
                pass

        title = Path(filename).stem
        parts = [f"# {title}\n"]

        for page_num, page in enumerate(doc, 1):
            text = page.get_text()
            if text.strip():
                parts.append(f"\n## Page {page_num}\n\n{text.strip()}\n")

        content_text = "\n".join(parts)
        page_count = len(doc)
        if not content_text.strip() or len(content_text.strip()) < 20:
            for page_num, page in enumerate(doc, 1):
                blocks = page.get_text("blocks")
                block_texts = []
                for b in blocks:
                    if b[4].strip():
                        block_texts.append(b[4].strip())
                if block_texts:
                    parts.append(f"\n## Page {page_num}\n\n" + "\n".join(block_texts) + "\n")
            content_text = "\n".join(parts)

        doc.close()

        return ParseResult(
            content=content_text,
            title=title,
            source_type="pdf",
            meta={
                "char_count": len(content_text),
                "page_count": page_count,
            },
        )
    except Exception as e:
        error_msg = str(e)
        if "no objects found" in error_msg.lower():
            raise ParseError(
                "Failed to parse PDF: file appears to be corrupted or is not a valid PDF. "
                "If this is a scanned/image-only PDF, OCR is required."
            )
        raise ParseError(f"Failed to parse PDF: {error_msg}")


# ── DOCX 解析器 ──

@register_parser("docx")
def parse_docx(content: bytes, filename: str) -> ParseResult:
    """DOCX 文档 → Markdown 纯文本。

    使用 python-docx 提取段落文本。
    """
    try:
        from io import BytesIO
        from docx import Document
    except ImportError:
        raise ParseError(
            "DOCX support requires python-docx. Install with: pip install python-docx"
        )

    if not content or len(content) < 100:
        raise ParseError("DOCX file is empty or too small")

    try:
        try:
            doc = Document(BytesIO(content))
        except Exception:
            import zipfile
            if not zipfile.is_zipfile(BytesIO(content)):
                raise ParseError(
                    "Failed to parse DOCX: file is not a valid DOCX document. "
                    "DOCX files are ZIP-based, the uploaded file may be corrupted or in old .doc format."
                )
            raise

        title = Path(filename).stem
        parts = [f"# {title}\n"]

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style else ""
            if style_name.startswith("Heading"):
                try:
                    level = int(style_name.replace("Heading", "").strip())
                    level = min(max(level, 1), 6)
                    parts.append(f"\n{'#' * level} {text}\n")
                except ValueError:
                    parts.append(f"\n{text}\n")
            elif style_name == "Title":
                parts.insert(0, f"# {text}\n")
            else:
                parts.append(f"{text}\n")

        for table in doc.tables:
            parts.append("\n")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append("| " + " | ".join(cells) + " |")
            parts.append("\n")

        content_text = "\n".join(parts)
        return ParseResult(
            content=content_text,
            title=title,
            source_type="docx",
            meta={
                "char_count": len(content_text),
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            },
        )
    except ParseError:
        raise
    except Exception as e:
        error_msg = str(e)
        if "not a zip file" in error_msg.lower():
            raise ParseError(
                "Failed to parse DOCX: file is not a valid DOCX document. "
                "If this is an old .doc format (Word 97-2003), please convert it to .docx first."
            )
        raise ParseError(f"Failed to parse DOCX: {error_msg}")
